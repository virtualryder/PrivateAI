"""
File loaders: convert uploaded files to plain text.

Supported formats: PDF (PyMuPDF), DOCX (python-docx), TXT, MD.
Returns a list of page/section strings (not one big blob) to preserve
document structure for downstream chunking.
"""

from pathlib import Path


def load_pdf(file_path: Path) -> list[str]:
    """Extract text from each page of a PDF. Returns list of page strings."""
    import fitz  # PyMuPDF
    pages = []
    with fitz.open(str(file_path)) as doc:
        for page in doc:
            text = page.get_text("text").strip()
            if text:
                pages.append(text)
    return pages


def load_docx(file_path: Path) -> list[str]:
    """
    Extract text from a DOCX file.
    Handles paragraphs and table cells (matches Job Search Tool pattern).
    Returns list of non-empty paragraph/cell strings.
    """
    from docx import Document
    doc = Document(str(file_path))
    sections = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            sections.append(text)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                sections.append(row_text)

    return sections


def load_text(file_path: Path) -> list[str]:
    """Load a plain text or markdown file. Returns list of non-empty lines."""
    text = file_path.read_text(encoding="utf-8", errors="replace")
    return [line.strip() for line in text.splitlines() if line.strip()]


_LOADERS = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".txt": load_text,
    ".md": load_text,
}

SUPPORTED_EXTENSIONS = list(_LOADERS.keys())


def load_file(file_path: Path) -> list[str]:
    """
    Dispatch to the correct loader based on file extension.
    Returns list of text sections.
    Raises ValueError for unsupported file types.
    """
    ext = file_path.suffix.lower()
    loader = _LOADERS.get(ext)
    if not loader:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {SUPPORTED_EXTENSIONS}")
    return loader(file_path)
